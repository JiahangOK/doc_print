import sqlite3
import xlrd
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
dic_num={
    1: '一',
    2: '二',
    3: '三',
    4: '四',
    5: '五',
    6: '六',
    7: '七',
    8: '八',
    9: '九',
    10: '十',
}
def createTables():
    conn = sqlite3.connect('wjh_yjb.db')
    print("Opened database successfully")
    c = conn.cursor()
    c.execute('CREATE TABLE YJB_RANK (ITEM TEXT NOT NULL,RANK INT NOT NULL,NUM INT NOT NULL);')
    print("Table YJB_RANK created successfully")
    c.execute('CREATE TABLE YJB_NUM_NAMES (ITEM TEXT NOT NULL,NUM INT NOT NULL,NAMES TEXT NOT NULL);')
    print("Table YJB_NUM_NAMES created successfully")
    conn.commit()
    conn.close()


def importData():
    conn = sqlite3.connect('wjh_yjb.db')
    print("Opened database successfully")
    c = conn.cursor()

    file_rank = "data/决赛名次.xlsx"
    file_num_name = "data/背番号.xlsx"

    book_rank = xlrd.open_workbook(file_rank)
    sheet_rank = book_rank.sheet_by_index(0)
    # 创建一个for循环迭代读取xls文件每行数据的
    for rx in range(0, sheet_rank.nrows):
        item = sheet_rank.cell(rx, 0).value
        rank = int(sheet_rank.cell(rx, 1).value)
        num = int(sheet_rank.cell(rx, 2).value)
        c.execute('INSERT INTO YJB_RANK (ITEM,RANK,NUM) VALUES (\'{}\', {}, {});'.format(item, rank, num))
    print("Table YJB_RANK inserted successfully")

    book_num_name = xlrd.open_workbook(file_num_name)
    sheet_num_name = book_num_name.sheet_by_index(0)
    # 创建一个for循环迭代读取xls文件每行数据的
    for rx in range(0, sheet_num_name.nrows):
        item = sheet_num_name.cell(rx, 0).value
        num = int(sheet_num_name.cell(rx, 1).value)
        names = sheet_num_name.cell(rx, 2).value
        c.execute('INSERT INTO YJB_NUM_NAMES (ITEM,NUM,NAMES) VALUES (\'{}\', {}, \'{}\');'.format(item, num, names))
    print("Table YJB_NUM_NAMES inserted successfully")

    conn.commit()
    conn.close()


def joinData():
    conn = sqlite3.connect('wjh_yjb.db')
    print("Opened database successfully")
    c = conn.cursor()
    c.execute('CREATE TABLE YJB_FINAL_TABLE (ITEM TEXT NOT NULL,RANK INT NOT NULL,NUM INT NOT NULL,NAMES TEXT NOT NULL);')
    conn.commit()

    c.execute('INSERT INTO YJB_FINAL_TABLE '
              'SELECT YJB_RANK.ITEM,YJB_RANK.RANK,YJB_RANK.NUM,YJB_NUM_NAMES.NAMES '
              'FROM YJB_RANK,YJB_NUM_NAMES '
              'WHERE YJB_RANK.NUM=YJB_NUM_NAMES.NUM AND YJB_RANK.ITEM=YJB_NUM_NAMES.ITEM;')
    conn.commit()

    conn.close()


def createWordDMT():
    conn = sqlite3.connect('wjh_yjb.db')
    print("Opened database successfully")
    c = conn.cursor()
    result = c.execute('SELECT * FROM YJB_FINAL_TABLE')

    document = Document()
    sections = document.sections
    for section in sections:
        # change orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(26)

    for row in result:
        item = row[0]
        rank = dic_num[row[1]]
        names = row[3]

        document.add_paragraph('')
        document.add_paragraph('')

        p = document.add_paragraph('')

        r = p.add_run('{} '.format(names))
        r.font.size = Pt(32)
        r.font.name = u'华文新魏'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')
        p.add_run('同学：'.format(names)).font.size = Pt(28)

        p = document.add_paragraph('')
        p.add_run('    荣获北京交通大学2019年学生体育“学院杯”').font.size = Pt(28)

        p = document.add_paragraph('')
        r = p.add_run('体育舞蹈')
        r.font.size = Pt(28)
        r.bold = True

        r = p.add_run(' 组')
        r.font.size = Pt(28)

        r = p.add_run(' {} '.format(item))
        r.bold = True
        r.font.size = Pt(28)

        p.add_run('项目').font.size = Pt(28)

        p = document.add_paragraph('')
        r = p.add_run('第 {} 名'.format(rank))
        r.bold = True
        r.font.size = Pt(42)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('')
        r = p.add_run('北京交通大学')
        r.font.size = Pt(28)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = document.add_paragraph('')
        r = p.add_run('2019 年 5 月')
        r.font.size = Pt(28)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        document.save('证书.docx')


if __name__ == '__main__':
    # 1.
    # createTables()

    # 2.
    # importData()

    # 3.
    # joinData()

    # 4.
    createWordDMT()
